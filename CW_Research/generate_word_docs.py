"""
ALGET CW Research — Word Document Generator
Generates all CW research documents as .docx files including:
- Core documents (01-06) with multi-agent review improvements
- Survey instruments (07-10)
** FULLY TRANSLATED TO ENGLISH **
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os

OUT_DIR = os.path.join(os.path.dirname(__file__), "word_output")
os.makedirs(OUT_DIR, exist_ok=True)

# ── Helpers ──────────────────────────────────────────────
def new_doc(title, subtitle=""):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(12)
        r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()  # spacer
    return doc

def h1(doc, text):
    doc.add_heading(text, level=1)

def h2(doc, text):
    doc.add_heading(text, level=2)

def h3(doc, text):
    doc.add_heading(text, level=3)

def para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p

def bullet(doc, text, level=0):
    doc.add_paragraph(text, style='List Bullet')

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

def doc_control(doc, fields):
    """Add document control table at top."""
    table = doc.add_table(rows=len(fields), cols=2)
    table.style = 'Light Grid Accent 1'
    for i, (k, v) in enumerate(fields):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for p in table.rows[i].cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
    doc.add_paragraph()

def save(doc, name):
    path = os.path.join(OUT_DIR, name)
    doc.save(path)
    print(f"  ✅ {name}")

def blank_line(doc, text="", n=1):
    for _ in range(n):
        p = doc.add_paragraph(text)
    return p

def checkbox_line(doc, text, checked=False):
    prefix = "☑ " if checked else "☐ "
    doc.add_paragraph(prefix + text)


# ═══════════════════════════════════════════════════════════
# 07 — Expert Background Survey 
# ═══════════════════════════════════════════════════════════
def gen_07():
    doc = new_doc("Expert Background Survey", "ALGET CW Pre-Evaluation Survey")
    
    doc_control(doc, [
        ("Evaluator ID", "E__ (E1–E5)"),
        ("Date", ""),
    ])

    h1(doc, "Section 1: Basic Information")
    para(doc, "1. Name: ___________________________")
    para(doc, "2. Affiliation: ___________________________")
    para(doc, "3. Title/Role: ___________________________")
    para(doc, "4. Highest Degree: ☐ Bachelor's  ☐ Master's  ☐ PhD  ☐ Other: ____")

    h1(doc, "Section 2: Areas of Expertise")
    para(doc, "5. Primary Expertise (Check all that apply):")
    checkbox_line(doc, "UX/HCI/Usability")
    checkbox_line(doc, "Learning Sciences / Instructional Design")
    checkbox_line(doc, "Mechanical Engineering (Statics/Dynamics)")
    checkbox_line(doc, "Educational Technology (EdTech)")
    checkbox_line(doc, "AI/ML Systems")
    checkbox_line(doc, "Other: ___________________________")

    para(doc, "6. Years of UX/Usability Evaluation Experience: ☐ None  ☐ 1-2 years  ☐ 3-5 years  ☐ 5+ years")
    para(doc, "7. Number of Cognitive Walkthroughs Conducted: ☐ None  ☐ 1-2 times  ☐ 3-5 times  ☐ 5+ times")
    para(doc, "8. e-Learning/EdTech Evaluation Experience: ☐ None  ☐ 1-2 times  ☐ 3-5 times  ☐ 5+ times")

    h1(doc, "Section 3: Domain Knowledge")
    para(doc, "9. Knowledge level in Statics / Dynamics:")
    para(doc, "   ☐ None  ☐ Beginner  ☐ Intermediate  ☐ Expert")
    para(doc, "10. Knowledge level in Bio-Inspired Design (Biomimicry):")
    para(doc, "   ☐ None  ☐ Beginner  ☐ Intermediate  ☐ Expert")

    h1(doc, "Section 4: AI Tool Experience")
    para(doc, "11. Experience using Generative AI tools (ChatGPT, Gemini, etc.):")
    para(doc, "   ☐ None  ☐ Rarely  ☐ Frequently  ☐ Daily")
    para(doc, "12. Experience with AI-powered learning tools:")
    para(doc, "   ☐ None  ☐ 1-2 tools  ☐ 3-5 tools  ☐ 5+ tools")
    para(doc, "13. Experience with AI Tutors / Intelligent Textbooks (List if any):")
    para(doc, "   ___________________________")

    h1(doc, "Section 5: Additional Information")
    para(doc, "14. Please list any specific perspectives or concerns you plan to focus on during this evaluation:")
    blank_line(doc, "", 3)

    save(doc, "07_Expert_Background_Survey.docx")

# ═══════════════════════════════════════════════════════════
# 08 — Post-Evaluation Reflection
# ═══════════════════════════════════════════════════════════
def gen_08():
    doc = new_doc("Post-Evaluation Reflection", "ALGET CW Post-Evaluation Survey")
    
    doc_control(doc, [
        ("Evaluator ID", "E__ (E1–E5)"),
        ("Date", ""),
        ("Total Evaluation Time", "_____ minutes"),
    ])

    h1(doc, "Section 1: Evaluation Process Feedback")
    para(doc, "1. How difficult was it to complete the CW Step Sheet?")
    para(doc, "   ☐ Very Easy  ☐ Easy  ☐ Neutral  ☐ Difficult  ☐ Very Difficult")
    para(doc, "2. Did the 4 defined Tasks adequately cover ALGET's core features?")
    para(doc, "   ☐ Strongly Agree  ☐ Agree  ☐ Neutral  ☐ Disagree  ☐ Strongly Disagree")
    para(doc, "3. If you feel any critical tasks were missing, please suggest them:")
    blank_line(doc, "", 2)
    para(doc, "4. How useful were the Learning Add-on questions (L1–L3 / L4–L5)?")
    para(doc, "   ☐ Very Useful  ☐ Useful  ☐ Neutral  ☐ Not Very Useful  ☐ Not at all Useful")

    h1(doc, "Section 2: Overall Impression of ALGET")
    para(doc, "5. How would you rate the overall pedagogical value of ALGET? (1=Very Low, 5=Very High)")
    para(doc, "   ☐ 1  ☐ 2  ☐ 3  ☐ 4  ☐ 5")
    para(doc, "6. Top 3 Strengths of ALGET:")
    para(doc, "   (1) ___________________________")
    para(doc, "   (2) ___________________________")
    para(doc, "   (3) ___________________________")
    para(doc, "7. Top 3 Weaknesses of ALGET:")
    para(doc, "   (1) ___________________________")
    para(doc, "   (2) ___________________________")
    para(doc, "   (3) ___________________________")

    h1(doc, "Section 3: AI Agent Evaluation")
    para(doc, "8. Rate the educational usefulness of each AI Agent (1=Very Low, 5=Very High):")
    add_table(doc,
        ["Agent", "1", "2", "3", "4", "5", "N/A"],
        [
            ["Orchestrator (Intent Routing)", "☐","☐","☐","☐","☐","☐"],
            ["Tutor (Socratic Scaffolding)", "☐","☐","☐","☐","☐","☐"],
            ["Janine (Life's Principles Validation)", "☐","☐","☐","☐","☐","☐"],
            ["Engineer (FBD/Simulation)", "☐","☐","☐","☐","☐","☐"],
        ])
    doc.add_paragraph()

    para(doc, "9. Overall quality of AI responses (Accuracy, Relevance, Helpfulness):")
    para(doc, "   ☐ 1  ☐ 2  ☐ 3  ☐ 4  ☐ 5")
    para(doc, "10. Did AI response latency negatively impact the learning experience?")
    para(doc, "   ☐ Not at all  ☐ Slightly  ☐ Moderately  ☐ Significantly  ☐ Very Much")

    h1(doc, "Section 4: Recommendations")
    para(doc, "11. The single most critical issue that must be fixed immediately:")
    blank_line(doc, "", 2)
    para(doc, "12. A long-term structural issue that needs redesign:")
    blank_line(doc, "", 2)
    para(doc, "13. Additional Comments:")
    blank_line(doc, "", 3)

    save(doc, "08_Post_Evaluation_Reflection.docx")

# ═══════════════════════════════════════════════════════════
# 09 — Consent Form
# ═══════════════════════════════════════════════════════════
def gen_09():
    doc = new_doc("Informed Consent Form", "ALGET Cognitive Walkthrough Participation Consent")

    h1(doc, "Study Information")
    para(doc, "Study Title: Evaluating the Learnability and Pedagogical Usability of the ALGET Intelligent Textbook Prototype")
    para(doc, "Researcher: ___________________________")
    para(doc, "Department: Educational Technology (Insert Course Name)")
    para(doc, "Contact: ___________________________")

    h1(doc, "Purpose of the Study")
    para(doc, "The purpose of this study is to evaluate the usability and educational effectiveness of the ALGET Intelligent Textbook prototype using the expert-based Cognitive Walkthrough methodology.")

    h1(doc, "What Participation Involves")
    bullet(doc, "Completing a pre-evaluation background survey (approx. 10 mins)")
    bullet(doc, "Conducting an independent Cognitive Walkthrough evaluation (approx. 90 mins)")
    bullet(doc, "Completing a post-evaluation reflection survey (approx. 10 mins)")
    bullet(doc, "Participating in a consolidation discussion session (approx. 90 mins, optional)")
    bullet(doc, "Total estimated time: 2–3.5 hours")

    h1(doc, "Risks and Benefits")
    para(doc, "There are minimal risks associated with participation. A potential benefit is gaining practical experience with advanced usability evaluation methods applied to GenAI educational systems.")

    h1(doc, "Confidentiality")
    para(doc, "All collected data will be used strictly for research/academic purposes. Your identity will be anonymized in any reports (e.g., labeled simply as 'Evaluator E1'). Raw data will be stored securely and destroyed after the completion of the research.")

    h1(doc, "Voluntary Participation")
    para(doc, "Your participation is completely voluntary. You may withdraw from the study at any time without any penalty and without providing a reason.")

    h1(doc, "Consent")
    para(doc, "I have read and understood the information provided above. I voluntarily agree to participate in this study.")
    doc.add_paragraph()
    para(doc, "Participant Signature: ___________________________ Date: ___________")
    para(doc, "Participant Name (Printed): ___________________________")
    doc.add_paragraph()
    para(doc, "Researcher Signature: ___________________________ Date: ___________")

    save(doc, "09_Consent_Form.docx")

# ═══════════════════════════════════════════════════════════
# 10 — Session Observation Sheet
# ═══════════════════════════════════════════════════════════
def gen_10():
    doc = new_doc("Session Observation Sheet", "ALGET CW Consolidation Session (Coordinator)")

    doc_control(doc, [
        ("Session Date", ""),
        ("Coordinator", ""),
        ("Participants", "E1 / E2 / E3 / E4 / E5"),
        ("Duration", "Start: _____ End: _____"),
    ])

    h1(doc, "Part A: Consolidation Discussion Log")
    para(doc, "Record the discussion points and consensus for each major issue.")
    add_table(doc,
        ["Issue ID", "Task.Step", "Issue", "#Agree", "#Disagree", "Final Sev", "Key Discussion Points"],
        [["","","","","","",""] for _ in range(10)])

    doc.add_paragraph()
    h1(doc, "Part B: Disagreement Resolution Log")
    para(doc, "Record discussions for items with high severity variance ($\ge$ 2):")
    add_table(doc,
        ["Issue ID", "E1", "E2", "E3", "E4", "E5", "Agreed Sev", "Resolution Reason"],
        [["","","","","","","",""] for _ in range(5)])

    doc.add_paragraph()
    h1(doc, "Part C: Final Priority Assignments")
    add_table(doc,
        ["Issue ID", "Avg Severity", "#Flagged (/5)", "Priority Score", "Quick Win / Structural", "Owner"],
        [["","","","","",""] for _ in range(10)])

    doc.add_paragraph()
    h1(doc, "Part D: Session Observation Notes")
    para(doc, "Recurring Patterns observed during discussion:")
    blank_line(doc, "", 3)
    para(doc, "Most debated topic/feature:")
    blank_line(doc, "", 2)
    para(doc, "Areas of strong consensus:")
    blank_line(doc, "", 2)
    para(doc, "Areas where consensus was difficult:")
    blank_line(doc, "", 2)
    para(doc, "Additional Notes:")
    blank_line(doc, "", 3)

    save(doc, "10_Session_Observation_Sheet.docx")

# ═══════════════════════════════════════════════════════════
# 03-Enhanced — Expert Step Sheet 
# ═══════════════════════════════════════════════════════════
def gen_03_enhanced():
    doc = new_doc("Expert Step Sheet (Enhanced)", "ALGET CW Expert Worksheet — Multi-Agent Review Final Version")

    doc_control(doc, [
        ("Study / Course", "Educational Technology – ALGET Intelligent Textbook"),
        ("Prototype", "(URL / Figma / Build)"),
        ("Version / Date", "v2.0 (English) / 2026-02-28"),
        ("Evaluator Name", ""),
        ("Evaluator ID", "E__ (E1–E5)"),
        ("Evaluator Type", "☐ UX/HCI  ☐ Learning Sciences  ☐ Domain SME  ☐ EdTech/AI  ☐ Target Learner"),
        ("Evaluation Date", ""),
        ("Duration", "Start: _____ End: _____ Total: _____ min"),
    ])

    # ─ A. Assumptions ─
    h1(doc, "A. Evaluation Assumptions")
    h2(doc, "A1. Target Learner Persona")
    add_table(doc,
        ["Attribute", "Profile"],
        [
            ["Learner Type", "Undergraduate Mechanical Engineering (1st–2nd year)"],
            ["Prior Knowledge", "High school physics/math. No Statics knowledge."],
            ["Digital Literacy", "Medium (LMS experience, limited AI tools)"],
            ["AI Tool Exp.", "Uses ChatGPT occasionally. Never used an 'AI Textbook' ★NEW"],
            ["Context", "Doing a homework assignment (Statics + Biomimicry)"],
            ["Motivation", "Mixed (Curiosity + Deadline pressure)"],
        ])

    doc.add_paragraph()
    h2(doc, "A2. ALGET Intelligent Features")
    checkbox_line(doc, "Hints/Scaffolding — Tutor Agent", True)
    checkbox_line(doc, "Generative AI Tutor Feedback — Multi-Agent", True)
    checkbox_line(doc, "Metacognitive Support — Janine's reflective feedback", True)
    checkbox_line(doc, "Interactive Simulation — Engineer Agent", True)
    checkbox_line(doc, "Adaptive Learning Pathways")
    checkbox_line(doc, "Diagnostic Placement")
    checkbox_line(doc, "Learning Analytics Dashboard")

    doc.add_paragraph()
    h2(doc, "A3. AI Response Protocol ★NEW")
    para(doc, "Evaluation Mode: ☐ Demo Mode (Cached responses)  ☐ Live Mode (Real-time Generation)")
    para(doc, "* Demo Mode is recommended to ensure all 5 experts evaluate identical outputs for comparability.", italic=True)
    para(doc, "* If using Live Mode, note it as a limitation and score the 'AI Output Quality' per task.", italic=True)

    # ─ B. Tasks ─
    h1(doc, "B. Task List Summary")
    add_table(doc,
        ["Task ID", "Task Name", "Core Action"],
        [
            ["T1", "First Session & Content Navigation", "Dashboard → Navigate to Statics 1.1"],
            ["T2", "Bio-Design Lab Inquiry & AI Interpretation", "Prompt input → Interpret Synthesis"],
            ["T3", "Interpreting Scaffolding & Validation", "Use Tutor question + Janine critique"],
            ["T4", "Operating Simulations & Verifying Learning", "Manipulate simulation → Concept link"],
        ])

    # ─ C. Severity ─
    doc.add_paragraph()
    h1(doc, "C. Severity Scale (0–4)")
    add_table(doc,
        ["Rating", "Label", "Criteria"],
        [
            ["0", "None", "Not a problem"],
            ["1", "Cosmetic", "Fix if there is time"],
            ["2", "Minor", "Slight frustration. Learning continues"],
            ["3", "Major", "Learning flow broken. Core feedback missed"],
            ["4", "Critical", "Learning impossible. Severe misconception risk"],
        ])

    # ─ D. Step Sheet Template (enhanced) ─
    doc.add_page_break()
    h1(doc, "D. Step Sheet — Walkthrough Tables")
    para(doc, "* Copy and paste the tables below for EACH Task (T1–T4).", bold=True)

    h2(doc, "[Task T__ : _________________________ ]")
    para(doc, "Scenario: ")
    para(doc, "Precondition: ")
    para(doc, "Success Criterion: ")

    h3(doc, "★ Step 0: First Impression Capture — NEW")
    para(doc, "Upon loading the first screen, observe for 3 seconds — What do you think this system is for?")
    para(doc, "Impression Memo: _____________________________________________")

    h3(doc, "Step-by-Step Walkthrough Table (Enhanced)")
    add_table(doc,
        ["Step#", "Learner Sub-goal", "Correct Action", "UI Cues Expected", "Likely Wrong Action ★"],
        [["","","","",""] for _ in range(6)])

    doc.add_paragraph()
    h3(doc, "CW Analysis + Learning Add-ons")
    add_table(doc,
        ["Step#", "Q1 Goal", "Q2 Find", "Q3 Understand", "Q4 Feedback", "L1 Objectives", "L2 Edu Feedback", "L3 Next Action"],
        [["","","","","","","",""] for _ in range(6)])

    doc.add_paragraph()
    para(doc, "★ L4–L5 (Optional for E2: Learning Sciences Expert):", bold=True)
    add_table(doc,
        ["Step#", "L4 Excessive Cognitive Load? (Y/N)", "L5 Promotes Metacognition? (Y/N)"],
        [["","",""] for _ in range(6)])

    doc.add_paragraph()
    h3(doc, "Issue Log")
    add_table(doc,
        ["Step#", "Issue Description", "Wrong Action/Interpretation", "Type", "Sev(0-4)", "Evidence", "Fix Idea"],
        [["","","","","","",""] for _ in range(5)])
    para(doc, "Type codes: Goal / Find / Understand / Feedback / Nav-Search / Learning")

    doc.add_paragraph()
    h3(doc, "★ AI Output Quality (Per Task, if using Live Mode) — NEW")
    add_table(doc,
        ["Criteria", "Score (1–5)", "Notes"],
        [
            ["Relevance", "", ""],
            ["Accuracy", "", ""],
            ["Helpfulness", "", ""],
            ["Latency", "☐Good ☐Slow ☐Very Slow", ""],
        ])

    doc.add_paragraph()
    h3(doc, "★ Content Accuracy (E3: Domain SME ONLY) — NEW")
    para(doc, "Factual correctness of AI-generated engineering content:")
    add_table(doc,
        ["Step#", "Content Element (e.g., FBD)", "Accurate? (Y/N)", "Error Description"],
        [["","","",""] for _ in range(3)])

    # ─ E. Task Summary ─
    doc.add_page_break()
    h1(doc, "E. Inter-Task Summary (Complete after each task)")
    h2(doc, "Task T__ Summary")
    add_table(doc,
        ["Item", "Details"],
        [
            ["Fragile Steps (Top 1–3)", "Step __: / Step __: / Step __:"],
            ["Dominant Breakdown", "☐Goal ☐Find ☐Understand ☐Feedback ☐Nav/Search ☐Learning"],
            ["Learnability Risk", "☐Low ☐Medium ☐High"],
            ["Priority Fixes (Sev 3–4)", "1. / 2. / 3."],
        ])

    # ─ F. Overall Summary ─
    doc.add_paragraph()
    h1(doc, "F. Overall Summary (Complete at end)")
    h2(doc, "F1. Recurring Problem Patterns")
    add_table(doc, ["Pattern", "Relevant Tasks", "Frequency"], [["","",""] for _ in range(3)])
    doc.add_paragraph()
    h2(doc, "F2. Most Critical Flaw in Learning Support")
    blank_line(doc, "", 3)
    h2(doc, "F3. Issues with AI/Intelligent Features")
    blank_line(doc, "", 3)
    h2(doc, "F4. Quick Wins vs. Structural Fixes")
    add_table(doc,
        ["Quick Wins", "Structural Fixes"],
        [["",""] for _ in range(3)])
    doc.add_paragraph()
    h2(doc, "F5. Additional Comments")
    blank_line(doc, "", 3)

    save(doc, "03_Expert_Step_Sheet_Enhanced.docx")

# ═══════════════════════════════════════════════════════════
# 05-Enhanced — Consolidation Table 
# ═══════════════════════════════════════════════════════════
def gen_05_consolidated():
    doc = new_doc("5-Expert Consolidation Table", "ALGET CW Master Results Aggregation")

    doc_control(doc, [
        ("Coordinator", ""),
        ("Consolidation Date", ""),
    ])

    h1(doc, "Master Consolidation Table")
    para(doc, "Priority Score = Avg Severity × #Experts Flagged")
    para(doc, "* Note: Recommend printing this table in Landscape mode.", italic=True)

    add_table(doc,
        ["ID", "Task", "Step", "Issue Statement", "E1", "E2", "E3", "E4", "E5",
         "#Flag", "AvgSev", "Priority", "Type", "Fix"],
        [["","","","","","","","","","","","","",""] for _ in range(15)])

    doc.add_paragraph()
    h1(doc, "Priority Scoring Legend")
    add_table(doc,
        ["Priority Score", "Class", "Action"],
        [
            ["0–4.9", "Low", "Backlog. Fix if time permits."],
            ["5.0–9.9", "Medium", "Fix in the next iteration."],
            ["10.0–14.9", "High", "Must fix in current iteration."],
            ["15.0–20.0", "Critical", "Fix immediately. Block release."],
        ])

    doc.add_paragraph()
    h1(doc, "Problem Type Distribution")
    add_table(doc,
        ["Type", "T1", "T2", "T3", "T4", "Total", "%"],
        [
            ["Goal","","","","","",""],
            ["Find","","","","","",""],
            ["Understand","","","","","",""],
            ["Feedback","","","","","",""],
            ["Nav/Search","","","","","",""],
            ["Learning","","","","","",""],
            ["Total","","","","","",""],
        ])

    save(doc, "05_Consolidation_Table.docx")

# ═══════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("Generating ALGET CW Word Documents (English)...")
    print(f"Output directory: {OUT_DIR}\n")
    gen_03_enhanced()
    gen_05_consolidated()
    gen_07()
    gen_08()
    gen_09()
    gen_10()
    print(f"\n✅ All English documents generated in: {OUT_DIR}")
