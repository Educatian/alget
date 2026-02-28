import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = docx.Document()
doc.add_heading('ALGET: System Architecture & Data Streamline', 0)

doc.add_heading('1. System Overview', level=1)
p = doc.add_paragraph('ALGET operates on a modern, decoupled architecture designed for high pedagogical fidelity and robust research data collection. The system is divided into four primary tiers:')
p.style.font.size = Pt(11)

doc.add_paragraph('• Client Interface (Frontend): A rich, interactive React application providing the structural textbook experience, formative assessments, and an Open Learner Model (OLM) Dashboard.', style='List Bullet')
doc.add_paragraph('• Intelligent Services (FastAPI Backend): A Python-based server that orchestrates a multi-agent framework utilizing Google Gemini models for real-time Socratic tutoring, contextual assessments, and narrative generation.', style='List Bullet')
doc.add_paragraph('• Data & Identity Infrastructure (Supabase): A managed PostgreSQL cloud database handling continuous telemetry streams, knowledge tracing records, and anonymous session authentication.', style='List Bullet')
doc.add_paragraph('• External Intelligence (Google GenAI): Integration with Gemini 2.0 architectures for generative pedagogical affordances.', style='List Bullet')

doc.add_heading('2. Core Subsystems', level=1)

doc.add_heading('A. The Structural Observability Engine', level=2)
p = doc.add_paragraph('The frontend acts as a passive observer extracting Structural Interactions:')
doc.add_paragraph('1. Time-on-Task (ToT): Calculated per section via React component lifecycles.', style='List Number')
doc.add_paragraph('2. Affordance Toggles: Records when users open or close supplementary learning modules.', style='List Number')
doc.add_paragraph('3. Active Metacognition: Tracks the volume/length of user highlights and annotations.', style='List Number')
doc.add_paragraph('4. Batch Dispatch: Events are buffered locally and dispatched to Supabase in batches.', style='List Number')

doc.add_heading('B. The Context-Aware Assessment & Knowledge Tracing System', level=2)
p = doc.add_paragraph('Unlike static quiz arrays, ALGET evaluation loop is dynamic:')
doc.add_paragraph('1. Generation: The client passes current context to the backend.', style='List Number')
doc.add_paragraph('2. Agentic Synthesis: The Assessment Agent queries Gemini to formulate questions.', style='List Number')
doc.add_paragraph('3. Local Knowledge Tracing: Calculates mastery using an Exponential Moving Average algorithm mapping correct/incorrect outcomes to a fluid 0.0 to 1.0 spectrum.', style='List Number')
doc.add_paragraph('4. Database Persistence: Mastery nodes are securely tied to the learner avoiding FK violations via invisible guest-auth.', style='List Number')

doc.add_heading('C. The Intervention Ecosystem (BigAL)', level=2)
doc.add_paragraph('If the learner goes idle for >90 seconds, or submits heavily penalized attempts, the system flags a "stuck_event" packet in Supabase. This triggers BookLayout to automatically reveal the Socratic BigAL intervention rail.')

doc.add_heading('D. The Open Learner Model (OLM)', level=2)
doc.add_paragraph('The Analytics Dashboard closes the system feedback loop. Real-time mastery scores, attempts, and calculated confidence levels converge into an intuitive layout reserved for research administration. The dashboard executes secure read-only queries against Supabase tables.')

doc.save('ALGET_System_Architecture.docx')
print('Word document saved as ALGET_System_Architecture.docx')
